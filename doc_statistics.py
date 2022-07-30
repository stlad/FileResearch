import docx, os
import zipfile as zip


def ger_docx_meta(docname): # вооще из дока можнов вытащить все что у годно. тут только минимум
    doc = docx.Document(docname)
    props = doc.core_properties
    info ={
        'Author':props.author,
        'Last modified by': props.last_modified_by,
        'Creation Date':f'{props.created.day}.{props.created.month}.{props.created.year}',
        'Last Mod Date':f'{props.modified.day}.{props.modified.month}.{props.modified.year}',
        'Last Print Date':f' {props.last_printed}',
        'Saves Count':props.revision
    }
    return info



class docx_staticstics:
    def __init__(self, name):
        self.name = name
        self.doc = docx.Document(self.name)
        self.paragraph_count = len(self.doc.paragraphs)
        self.fullpath = os.path.abspath(name)



    def get_info(self):
        paragraphs = self.doc
        for paragraph in self.doc.paragraphs:
            print(paragraph.text)


    def get_par_from_index(self, index):
        p_list = list(self.doc.paragraphs)
        p = Paragraph(p_list[index-1])
        return p

    def get_doxc_zip_info(self):
        z = zip.ZipFile(self.fullpath, 'r')
        a = z.infolist()
        res = []
        for info in a:
            inf = {}
            inf['FileName'] = info.filename
            inf['CRC'] = info.CRC
            inf['Create system'] = info.create_system
            inf['Create version'] = info.create_version
            inf['Compress type'] = info.compress_type
            inf['File Size'] = info.file_size
            inf['Volume'] = info.volume
            res.append(inf)
        z.close()
        return res


class Paragraph:
    def __init__(self, par):
        self.paragraph = par
        self.text = par.text
        self.symbol_count = len(self.text)
        self.words_count = len(self.text.split(' '))
        self.style =self.paragraph.style

    def get_par_info(self):
        res ={}
        res['Text']= self.text,
        res['Symbol count']= str(self.symbol_count)
        res['Words Count']=str(self.words_count)
        res['Paragraph style']=str(self.style)
        res['Alignment '] = str(self.paragraph.alignment)
        return res




'''os.chdir(os.getcwd()+'/Files')

a = docx_staticstics('Document.docx')
print(a.paragraph_count)
'''